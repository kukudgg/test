import multiprocessing
import os
import threading
import time

from docx import Document
from docx.shared import Pt
from pdf2docx import Converter
from tqdm import tqdm

folderDirs = []
threads = []


# 转换PDF到docx
def add(pdf_dir, output_dir, filename):
    start_time = time.perf_counter();
    pdf_path = os.path.join(pdf_dir, filename)
    # 获取文件名（不包括扩展名）作为DOCX文件的名称
    docx_name = os.path.splitext(filename)[0] + '.docx'
    print(f"开始转换:{filename}")
    docx_path = os.path.join(output_dir, docx_name)
    # 调用转换函数
    pdf_to_docx(pdf_path, docx_path, Pt(8), 'Malgun Gothic')
    end_time = time.perf_counter();
    use_time = '%.2f' % (end_time - start_time)
    print(f"{docx_name}转换完成,用了{use_time}s")


# 转换PDF到docx
def pdf_to_docx(pdf_path, docx_path, font_size, font_name):
    # 创建Converter对象
    conv = Converter(pdf_path)
    # 开始转换
    conv.convert(docx_path, start=0, end=None)
    # 关闭Converter对象
    conv.close()
    # 打开转换后的DOCX文件，并设置字体大小和字体样式
    doc = Document(docx_path)
    for table in tqdm(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = font_size
    doc.save(docx_path)


# 创建多线程
def batch_convert(pdf_dir, output_dir):
    pdfs = []
    multiprocess = []
    # 遍历目录中的所有PDF文件
    for filename in os.listdir(pdf_dir):
        if filename.endswith('.pdf'):
            pdfs.append(filename)
    for pdf in pdfs:
        multiprocess.append(
            multiprocessing.Process(target=add, args=(pdf_dir, output_dir, pdf))
        )
    for p in multiprocess:
        time.sleep(0.1)
        p.start()
    for p in multiprocess:
        time.sleep(0.1)
        p.join()


# 创建文件夹
def create_folders(source_folder, target_folder):
    print("正在创建文件夹")
    # 遍历源文件夹中的所有文件夹
    for root, dirs, files in os.walk(source_folder):
        # 在目标文件夹中创建相应的文件夹
        for dir_name in dirs:
            folder_list = []
            folder_list.append(f'{source_folder}\\{dir_name}')
            folder_list.append(f'{target_folder}\\{dir_name}')
            folderDirs.append(folder_list)
            target_dir = os.path.join(target_folder, dir_name)
            os.makedirs(target_dir, exist_ok=True)
    print("创建成功")


def fold_super(source_folder, target_folder):
    # 创建文件夹
    create_folders(source_folder, target_folder)
    # input_path = r'E:\demo_pdf\000001_平安银行'
    # output_path = r'E:\demo_word\000001_平安银行'
    for item in folderDirs:
        threads.append(threading.Thread(batch_convert(item[0], item[1])))
    for th in threads:
        time.sleep(0.1)
        th.start()
    for th in threads:
        time.sleep(0.1)
        th.join()


if __name__ == '__main__':
    # 源文件夹路径
    source_folder = './pdf'
    # 目标文件夹路径
    target_folder = './docx'
    fold_super(source_folder, target_folder)
    print("success")
